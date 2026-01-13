# =========================================
# Tractive Effort Tool DOCX Report Builder Module
# =========================================
# Professional document generation for tractive effort engineering reports.
#
# Purpose: Create formatted Microsoft Word documents containing complete
# tractive effort analysis reports with resistance components, power calculations,
# and electrical requirements for railway engineering documentation and procurement.
#
# Engineering Context:
# - Generates professional engineering documentation for traction analysis
# - Ensures consistent formatting across all railway power reports
# - Supports railway industry standards for locomotive specification
# - Provides timestamped, auditable traction requirement records
#
# Used by: api.py (/download-report endpoint)
# Dependencies: python-docx (document generation), service.py (report formatting)

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

def create_te_docx_report(inputs: Dict[str, Any], results: Dict[str, Any], inputs_raw: Dict[str, Any] = None) -> io.BytesIO:
    """
    Create a professional DOCX report for tractive effort and power analysis.

    Generates a formatted Microsoft Word document containing the complete
    engineering analysis report with input parameters, tractive effort calculations,
    resistance component breakdown, and electrical power requirements.

    Document Structure:
    1. Title and timestamp header
    2. Input parameters section with operating conditions
    3. Calculation results summary with key metrics
    4. Resistance components breakdown with engineering explanations

    Args:
        inputs (Dict[str, Any]): Complete input parameters from validation layer
                                containing load, weights, gradients, speeds, and modes

        results (Dict[str, Any]): Complete calculation results from service layer
                                 containing tractive effort, power, current, and resistance components

        inputs_raw (Dict[str, Any], optional): Original user inputs for display
                                             - Falls back to processed inputs if not provided
                                             - Preserves user-entered values in reports

    Returns:
        io.BytesIO: In-memory byte stream containing the complete DOCX file
                   - Ready for HTTP response streaming
                   - Compatible with FastAPI StreamingResponse
                   - No temporary files created on disk

    Document Sections:

    HEADER:
    - Report title: "Tractive Effort Calculation Report"
    - Generation timestamp for version control and audit trails

    INPUT PARAMETERS:
    - Train load (shunting load) in tonnes
    - Locomotive gross weight (GBW) in tonnes
    - Track gradient with specification method (degrees or 1 in G)
    - Track curvature with specification method (radius or degrees)
    - Operating speed in km/h
    - Operating mode (Start or Running)

    CALCULATION RESULTS SUMMARY:
    - Total tractive effort (TE) in kg and tonnes
    - Rail horsepower at wheel-rail interface
    - Overhead electric (OHE) current requirement in amperes

    RESISTANCE COMPONENTS BREAKDOWN:
    - T1: Wagon rolling resistance (load-dependent)
    - T2: Locomotive rolling resistance (weight-dependent)
    - T3: Gradient resistance (slope-dependent)
    - T4: Curvature resistance (curve geometry-dependent)

    Formatting Features:
    - Professional document styling with structured headings
    - Bold formatting for section headers and key results
    - Clear units and engineering terminology
    - Comprehensive resistance component analysis

    Engineering Standards Compliance:
    - Follows railway engineering documentation practices
    - Includes all necessary traction calculation parameters
    - Provides traceable resistance component analysis
    - Suitable for locomotive procurement and electrification planning

    Technical Notes:
    - Uses in-memory streams to avoid disk I/O in cloud environments
    - Compatible with cloud deployment and containerized applications
    - Thread-safe document generation for concurrent requests
    - Minimal memory footprint for large-scale operations

    Error Handling:
    - Graceful handling of missing input data with fallbacks
    - Clear error messages for document generation failures
    - Fallback input handling maintains report completeness

    Example:
        >>> inputs = {'load': 1500.0, 'loco_weight': 120.0, 'gradient': 1.0,
        ...          'curvature': 300.0, 'speed': 60.0, 'mode': 'Running',
        ...          'grad_type': 'Degree', 'curvature_unit': 'Radius(m)'}
        >>> results = {'te': 15420.5, 'power': 245.8, 'ohe_current': 456.2,
        ...           'T1': 2025.3, 'T2': 350.6, 'T3': 1000.0, 'T4': 45.2}
        >>> docx_stream = create_te_docx_report(inputs, results)
        >>> # Returns BytesIO stream ready for download

    Note:
        The report provides comprehensive traction analysis for railway
        engineering applications, including both mechanical tractive effort
        and electrical power system requirements for complete system design.
    """

    # Handle optional raw inputs with fallback to processed inputs
    if inputs_raw is None:
        inputs_raw = inputs

    # Create new Word document with professional formatting
    doc = Document()

    # Add report header with title and timestamp
    doc.add_heading("Tractive Effort Calculation Report", 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Input Parameters Section
    doc.add_heading('1. Inputs', level=1)
    p = doc.add_paragraph()
    p.add_run(f"• Shunting Load: {inputs_raw.get('load', inputs.get('load'))} tons\n")
    p.add_run(f"• GBW of Vehicle: {inputs_raw.get('loco_weight', inputs.get('loco_weight'))} tons\n")
    p.add_run(f"• Gradient: {inputs_raw.get('gradient', inputs.get('gradient'))} ({inputs.get('grad_type')})\n")
    p.add_run(f"• Curvature: {inputs_raw.get('curvature', inputs.get('curvature'))} ({inputs.get('curvature_unit')})\n")
    p.add_run(f"• Speed: {inputs_raw.get('speed', inputs.get('speed'))} km/h\n")
    p.add_run(f"• Mode: {inputs.get('mode')}\n")

    # Calculation Results Summary Section
    doc.add_heading('2. Calculation Results', level=1)

    # Summary results with bold formatting
    p = doc.add_paragraph()
    run = p.add_run("Summary of Results:\n")
    run.bold = True
    p.add_run(f"  • Tractive Effort (TE): {results['te']:.2f} kg  ({results['te']/1000:.3f} tons)\n")
    p.add_run(f"  • Rail Horsepower: {results['power']:.2f} HP\n")
    p.add_run(f"  • OHE Current: {results['ohe_current']:.2f} A\n")

    # Resistance components breakdown with bold formatting
    p = doc.add_paragraph()
    run = p.add_run("\nResistance Components:\n")
    run.bold = True
    p.add_run(f"  • T1 (Wagon Rolling Resistance): {results['T1']:.2f} kg\n")
    p.add_run(f"  • T2 (Loco Rolling Resistance): {results['T2']:.2f} kg\n")
    p.add_run(f"  • T3 (Gradient Resistance): {results['T3']:.2f} kg\n")
    p.add_run(f"  • T4 (Curvature Resistance): {results['T4']:.2f} kg\n")

    try:
        # Save document to in-memory stream for efficient delivery
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)  # Reset stream position for reading

        return file_stream

    except Exception as e:
        # Provide clear error information for document generation failures
        raise RuntimeError(f"An error occurred while creating the docx report: {e}")