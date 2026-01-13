# =========================================
# Load Distribution Tool DOCX Report Builder Module
# =========================================
# Professional document generation for load distribution engineering reports.
#
# Purpose: Create formatted Microsoft Word documents containing complete
# load distribution analysis reports with safety assessments, suitable for
# regulatory submissions, design reviews, and professional railway engineering records.
#
# Engineering Context:
# - Generates professional engineering documentation for load analysis
# - Ensures consistent formatting across all safety reports
# - Supports railway industry standards for vehicle certification
# - Provides timestamped, auditable safety assessment records
#
# Used by: api.py (/download-report endpoint)
# Dependencies: python-docx (document generation), service.py (step formatting)

import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

def create_load_distro_docx(inputs: Dict[str, Any], results: Dict[str, Any]) -> io.BytesIO:
    """
    Create a professional DOCX report for load distribution safety analysis.

    Generates a formatted Microsoft Word document containing the complete
    engineering analysis report with input parameters, safety assessment,
    individual wheel loads, and detailed calculation methodology.

    Document Structure:
    1. Title and timestamp header
    2. Input parameters section with engineering context
    3. Calculation results summary with safety status and compliance
    4. Visual diagram (if available) for load distribution representation
    5. Detailed step-by-step calculation methodology

    Args:
        inputs (Dict[str, Any]): Complete input parameters from validation layer
                                containing config_type, total_load, percentages

        results (Dict[str, Any]): Complete analysis results from service layer
                                 containing wheel loads, safety ratios, and status

    Returns:
        io.BytesIO: In-memory byte stream containing the complete DOCX file
                   - Ready for HTTP response streaming
                   - Compatible with FastAPI StreamingResponse
                   - No temporary files created on disk

    Document Sections:

    HEADER:
    - Report title: "Load Distribution Safety Report"
    - Generation timestamp for version control and audit trails

    INPUT PARAMETERS:
    - Vehicle configuration type (Bogie/Truck/Axle)
    - Total vehicle load with units (tonnes)
    - Front axle load percentage
    - Individual wheel load percentages (Q1, Q3)

    CALCULATION RESULTS SUMMARY:
    - Overall safety status (SAFE/UNSAFE) with clear highlighting
    - ΔQ/Q ratio analysis with allowable limits
    - Front and rear axle loads
    - Individual wheel loads (Q1-Q4) in a clear format

    VISUAL DIAGRAM:
    - Attempts to include Diagram.png from project root
    - Provides visual representation of load distribution
    - Graceful fallback if diagram is not available

    DETAILED CALCULATION STEPS:
    - Complete step-by-step engineering methodology
    - Mathematical formulations and intermediate values
    - Safety limit validation process

    Formatting Features:
    - Professional document styling with structured headings
    - Table format for results summary with proper column sizing
    - Clear status indicators for safety compliance
    - Embedded diagram support for visual clarity

    Safety Assessment Display:
    - Status: "SAFE" (green/indicator) or "UNSAFE" (warning indicator)
    - ΔQ/Q Ratio: Actual load variation as percentage
    - Allowed Limit: Configuration-specific threshold (Bogie: 60%, Truck: 50%)
    - Individual wheel loads for detailed analysis

    Technical Notes:
    - Uses in-memory streams to avoid disk I/O in cloud environments
    - Compatible with cloud deployment and containerized applications
    - Thread-safe document generation for concurrent requests
    - Minimal memory footprint for large-scale operations

    Engineering Standards Compliance:
    - Follows railway engineering documentation practices
    - Includes all necessary safety assessment parameters
    - Provides traceable calculation methodology
    - Suitable for regulatory review and vehicle certification

    Error Handling:
    - Graceful handling of missing diagram files
    - Clear error messages for document generation failures
    - Fallback behavior maintains report completeness

    Example:
        >>> inputs = {'config_type': 'Bogie', 'total_load': 85.5, 'front_percent': 55.0,
        ...          'q1_percent': 52.0, 'q3_percent': 48.0}
        >>> results = {'status': 'SAFE', 'delta_q_by_q': 0.128, 'limit': 0.6,
        ...           'front_load': 46.525, 'rear_load': 38.975,
        ...           'q_values': {'Q1': 24.194, 'Q2': 22.331, 'Q3': 18.734, 'Q4': 20.241}}
        >>> docx_stream = create_load_distro_docx(inputs, results)
        >>> # Returns BytesIO stream ready for download

    Note:
        The report automatically includes visual diagrams when Diagram.png
        exists in the project root directory, enhancing the engineering
        documentation with graphical load distribution representation.
    """

    # Determine diagram image path for visual load distribution representation
    current_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    image_path = os.path.join(current_dir, "Diagram.png")

    # Create new Word document with professional formatting
    doc = Document()

    # Add report header with title and timestamp
    doc.add_heading('Load Distribution Safety Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Input Parameters Section
    doc.add_heading('1. Input Parameters', level=1)
    p = doc.add_paragraph()
    p.add_run(f"  • Configuration Type: {inputs['config_type']}\n")
    p.add_run(f"  • Total Load: {inputs['total_load']:.2f} Ton\n")
    p.add_run(f"  • Front Load Percentage: {inputs['front_percent']:.2f}%\n")
    p.add_run(f"  • Q1 Percentage (of Front Load): {inputs['q1_percent']:.2f}%\n")
    p.add_run(f"  • Q3 Percentage (of Rear Load): {inputs['q3_percent']:.2f}%\n")

    # Calculation Results Summary Section with professional table formatting
    doc.add_heading('2. Calculation Results Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(4.0)  # Wider column for detailed results
    table.columns[1].width = Inches(2.5)  # Narrower column for diagram

    # Populate results summary with safety status and load distribution details
    cell_text = (
        f"Overall Status: {results['status'].upper()}\n\n"
        f"ΔQ/Q Ratio: {results['delta_q_by_q']:.2%}\n"
        f"Allowed Limit: {results['limit']:.0%}\n\n"
        f"Front Load: {results['front_load']:.2f} Ton\n"
        f"Rear Load: {results['rear_load']:.2f} Ton\n\n"
        f"Q1: {results['q_values']['Q1']:.2f} Ton | Q2: {results['q_values']['Q2']:.2f} Ton\n"
        f"Q3: {results['q_values']['Q3']:.2f} Ton | Q4: {results['q_values']['Q4']:.2f} Ton"
    )
    table.cell(0, 0).text = cell_text

    # Handle diagram inclusion with graceful fallback
    p = table.cell(0, 1).paragraphs[0]
    if os.path.exists(image_path):
        try:
            # Attempt to embed diagram for visual load distribution representation
            p.add_run().add_picture(image_path, width=Inches(2.5))
        except Exception as e:
            # Fallback if diagram cannot be processed
            p.text = f"Diagram.png found but could not be added. Error: {e}"
    else:
        # Clear indication when diagram is not available
        p.text = "Diagram.png not found."

    # Detailed Calculation Steps Section
    doc.add_heading('3. Detailed Calculation Steps', level=1)
    detailed_steps_text = format_load_distro_steps(inputs, results)
    doc.add_paragraph(detailed_steps_text)

    try:
        # Save document to in-memory stream for efficient delivery
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)  # Reset stream position for reading

        return file_stream

    except Exception as e:
        # Provide clear error information for document generation failures
        raise RuntimeError(f"An error occurred while creating the report: {e}")

def format_load_distro_steps(inputs: Dict[str, Any], results: Dict[str, Any]) -> str:
    """
    Format load distribution calculation steps (compatibility function).

    This function provides backward compatibility by importing and calling
    the format_load_distro_steps function from the service module. It ensures
    consistent report formatting across different report generation entry points.

    Args:
        inputs (Dict[str, Any]): Validated input parameters for the analysis
        results (Dict[str, Any]): Complete calculation results and safety assessment

    Returns:
        str: Formatted calculation report string with detailed methodology

    Note:
        This is a compatibility wrapper - the actual formatting logic
        is implemented in the service module for better separation of concerns
        and consistent formatting across API and document generation.
    """
    from ..service import format_load_distro_steps as service_format
    return service_format(inputs, results)