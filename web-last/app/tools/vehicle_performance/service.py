# =========================================
# Vehicle Performance Tool Service Layer
# =========================================
# Orchestrates vehicle performance calculations and generates reports.
# Contains physics-based resistance calculations, performance analysis,
# and document generation for locomotive performance evaluation.
#
# Purpose: Provide comprehensive vehicle performance analysis including
#          traction limits, speed calculations, and performance curves.
# Used by: vehicle_performance API endpoints and report generation
# Dependencies: numpy (for interpolation), python-docx (for reports)

import math  # For trigonometric and mathematical functions
import io  # For in-memory file streams
from datetime import datetime  # For report timestamps
from typing import Dict, List, Any  # Type hints for better code documentation

# --- Try to import required libraries ---
try:
    import numpy as np  # Required for torque curve interpolation
except ImportError:
    print("="*50)
    print("ERROR: 'numpy' library not found for VehiclePerformance.")
    print("Install it using: pip install numpy")
    print("="*50)
    np = None

try:
    import docx  # Required for .docx report generation
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("="*50)
    print("ERROR: 'python-docx' library not found.")
    print("Install it using: pip install python-docx")
    print("="*50)
    docx = None

# ===============================================================


# Core Physics Functions
# (Extracted from vehicle_performance_calculator.py)
# Units:
# - Speed: km/h
# - Weight: metric tons
# - Forces: Newtons (N)
# - Torque: Newton-meters (Nm), RPM in rev/min, Power in kW
# - Slope: percent (%), curve: degree
# ===============================================================

def rolling_resistance_loco(speed_kmh: float, weight_ton: float, num_axles: int) -> float:
    """
    Calculate rolling resistance for locomotive in Newtons.

    Rolling resistance is the force required to overcome bearing friction,
    wheel deformation, and track irregularities for locomotives.

    Args:
        speed_kmh (float): Locomotive speed in km/h
        weight_ton (float): Locomotive weight in metric tons
        num_axles (int): Number of locomotive axles

    Returns:
        float: Rolling resistance force in Newtons

    Formula:
        F_rr = (A + B×V + C×V²) × W × g
        Where:
        - A, B, C are empirical coefficients
        - V is speed in km/h
        - W is weight in tons
        - g is gravity (9.81 m/s²)

    Physics:
        Coefficients account for different resistance characteristics
        of locomotive wheels vs wagon wheels
    """
    if weight_ton <= 0 or num_axles <= 0:
        return 0.0
    A = 0.647 + (13.17 / (weight_ton / num_axles))  # Base resistance coefficient
    B = 0.00933  # Speed-dependent coefficient
    C = 0.057 / weight_ton  # Speed-squared coefficient
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def rolling_resistance_wagon(speed_kmh: float, weight_ton: float) -> float:
    """
    Calculate rolling resistance for wagon in Newtons.

    Rolling resistance for freight wagons, typically lower than locomotives
    due to different wheel designs and axle configurations.

    Args:
        speed_kmh (float): Wagon speed in km/h
        weight_ton (float): Wagon weight in metric tons

    Returns:
        float: Rolling resistance force in Newtons

    Formula:
        F_rr = (A + B×V + C×V²) × W × g
        Where coefficients are empirically derived for wagon wheels

    Note:
        Wagon resistance is generally lower than locomotive resistance
        due to optimized wheel profiles and bearing designs
    """
    if weight_ton <= 0:
        return 0.0
    A = 0.6438797  # Base resistance coefficient for wagons
    B = 0.01047218  # Speed coefficient
    C = 0.00007323  # Speed-squared coefficient
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def gradient_resistance(weight_ton: float, slope_percent: float) -> float:
    """
    Calculate gradient (slope) resistance in Newtons.

    The additional force required to move uphill against gravity.
    This is the component of gravitational force parallel to the slope.

    Args:
        weight_ton (float): Vehicle weight in metric tons
        slope_percent (float): Slope gradient as percentage (e.g., 2.0 = 2%)

    Returns:
        float: Gradient resistance force in Newtons

    Physics:
        F_gradient = m × g × (slope/100)
        Where slope/100 converts percentage to decimal fraction

    Example:
        100 ton vehicle on 2% slope: F = 100 × 9.81 × 0.02 = 196.2 N
    """
    return weight_ton * 1000 * 9.81 * slope_percent / 100.0

def curvature_resistance(weight_ton: float, curve_degree: float) -> float:
    """
    Calculate curvature (curve) resistance in Newtons.

    Additional resistance when negotiating curved track due to:
    - Longer path traveled by outer wheels
    - Flange friction on inner wheels
    - Cant deficiency effects

    Args:
        weight_ton (float): Vehicle weight in metric tons
        curve_degree (float): Curve sharpness in degrees

    Returns:
        float: Curvature resistance force in Newtons

    Formula:
        F_curve = 0.4 × W × D × g
        Where D is curve degree (empirical relationship)

    Note:
        The 0.4 coefficient is empirically derived for railway curves
        Resistance increases with tighter curves (higher degree values)
    """
    return 0.4 * weight_ton * curve_degree * 9.81

def starting_resistance_loco(weight_ton: float) -> float:
    """
    Calculate starting resistance for locomotive in Newtons.

    Additional resistance when starting from rest due to:
    - Static friction in bearings and wheels
    - Track adhesion requirements
    - Brake drag and mechanical lash

    Args:
        weight_ton (float): Locomotive weight in metric tons

    Returns:
        float: Starting resistance force in Newtons

    Formula:
        F_start = 6.0 × W × g
        Where 6.0 is an empirical coefficient for locomotives

    Note:
        Starting resistance is higher than running resistance
        due to static friction vs dynamic friction
    """
    return 6.0 * weight_ton * 9.81

def starting_resistance_wagon(weight_ton: float) -> float:
    """
    Calculate starting resistance for wagon in Newtons.

    Similar to locomotive but typically lower coefficient due to
    different wheel designs and fewer mechanical components.

    Args:
        weight_ton (float): Wagon weight in metric tons

    Returns:
        float: Starting resistance force in Newtons

    Formula:
        F_start = 4.0 × W × g
        Where 4.0 is an empirical coefficient for wagons

    Note:
        Wagon starting resistance is lower than locomotive
        due to simpler mechanical design
    """
    return 4.0 * weight_ton * 9.81

def interpolate_torque(engine_rpm: float, curve: Dict[int, float]) -> float:
    """
    Interpolate engine torque at a given RPM from discrete torque curve.

    Since torque curves are typically provided as discrete RPM-torque pairs,
    this function interpolates to find torque at any RPM value.

    Args:
        engine_rpm (float): Engine speed in RPM
        curve (Dict[int, float]): Torque curve as {rpm: torque_nm}

    Returns:
        float: Interpolated torque in Newton-meters

    Method:
        Uses numpy's linear interpolation between data points
        Extrapolates using endpoint values if outside range

    Raises:
        ValueError: If torque curve is empty
        ImportError: If numpy is not available

    Note:
        Linear interpolation assumes smooth torque curve
        Real engines may have more complex characteristics
    """
    if not curve:
        raise ValueError("Torque curve is empty.")
    if np is None:
        raise ImportError("Numpy is required for interpolation.")

    # Ensure keys are sorted for proper interpolation
    rpms = np.array(sorted(curve.keys()), dtype=float)
    torques = np.array([curve[r] for r in rpms], dtype=float)

    # np.interp handles extrapolation (clamps to endpoints)
    return float(np.interp(engine_rpm, rpms, torques))


# ===============================================================


class VehiclePerformanceCalculator:
    """
    Main calculator class for vehicle performance analysis.

    This class encapsulates all the business logic for calculating locomotive
    performance characteristics including traction limits, speed calculations,
    and performance curves. Adapted from the original Tkinter application.

    Key Features:
    - Traction vs slipping analysis
    - Speed vs slope calculations
    - Performance curve generation
    - DOCX report generation

    Architecture:
    - Pure calculation methods (no UI dependencies)
    - Comprehensive input validation
    - Modular design for different analysis types
    """

    def __init__(self, inputs: dict):
        """
        Initialize calculator with validated input parameters.

        Sets up all vehicle and track parameters for calculations.
        Performs unit conversions and parameter validation.

        Args:
            inputs (dict): Validated input parameters from API/validation layer

        Key Parameters:
        - Track: max_curve, curve_unit, max_slope, slope_unit
        - Vehicle: loco_gvw_kg, max_speed_kmh, num_axles, rear_axle_ratio
        - Powertrain: gear_ratios, peak_power_kw, torque_curve
        - Other: friction_mu, wheel_dia_m, min_rpm, max_rpm, shunting_load_t

        Unit Conversions:
        - Curve: meters to degrees if needed
        - Slope: degrees to percent if needed
        - Weight: kg to tons for calculations
        """
        self.inputs = inputs

        # --- Store key parameters from the validated inputs dict ---
        # Track parameters
        self.max_curve_deg = inputs.get('max_curve', 0.0)
        if inputs.get('curve_unit') == 'm' and self.max_curve_deg != 0:
            # Convert curve radius to equivalent degree
            self.max_curve_deg = 1750.0 / self.max_curve_deg

        self.max_slope_percent = inputs.get('max_slope', 0.0)
        if inputs.get('slope_unit') == 'degree':
            # Convert slope from degrees to percent
            self.max_slope_percent = math.tan(math.radians(self.max_slope_percent)) * 100.0

        # Vehicle parameters
        self.loco_gvw_kg = inputs.get('loco_gvw_kg', 0.0)
        self.loco_gvw_ton = self.loco_gvw_kg / 1000.0  # Convert to tons
        self.max_speed_kmh = inputs.get('max_speed_kmh', 0.0)
        self.num_axles = inputs.get('num_axles', 1)
        self.rear_axle_ratio = inputs.get('rear_axle_ratio', 1.0)
        self.gear_ratios = inputs.get('gear_ratios', [1.0])
        self.shunting_load_t = inputs.get('shunting_load_t', 0.0)

        # Powertrain parameters
        self.peak_power_kw = inputs.get('peak_power_kw', 0.0)
        self.friction_mu = inputs.get('friction_mu', 0.3)
        self.wheel_dia_m = inputs.get('wheel_dia_m', 0.73)
        self.min_rpm = inputs.get('min_rpm', 100)
        self.max_rpm = inputs.get('max_rpm', 2500)

        # Torque curve validation
        self.torque_curve = inputs.get('torque_curve', {})
        if not self.torque_curve:
            raise ValueError("Torque Curve is empty or missing.")

        self.max_torque_nm = max(self.torque_curve.values()) if self.torque_curve else 0.0

    def run_tractive_calculation(self) -> dict:
        """
        Calculate traction limits and slipping analysis.

        Determines the maximum tractive force the locomotive can produce
        and compares it to the force limit before wheels slip.

        Returns:
            dict: Traction analysis results containing:
                - max_traction_generated_n: Maximum force from powertrain (N)
                - max_traction_slipping_n: Maximum force before slipping (N)
                - result_message: "Limited by slipping" or "Not limited by slipping"

        Physics:
        - Generated traction: Based on max torque, gear ratios, wheel diameter
        - Slipping limit: μ × weight × g (friction coefficient limit)
        - Actual traction is minimum of generated and slipping limits

        Note:
            This calculation assumes all axles are driving axles
        """
        if self.wheel_dia_m == 0:
             raise ValueError("Wheel Diameter cannot be zero.")

        # Maximum traction from powertrain (using highest gear ratio)
        max_traction_generated_n = 2 * (self.max_torque_nm * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m

        # Maximum traction before slipping (friction limit)
        max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81

        # Determine limiting factor
        if max_traction_generated_n > max_traction_slipping_n:
            result_message = "Limited by slipping"
        else:
            result_message = "Not limited by slipping"

        return {
            "max_traction_generated_n": max(max_traction_generated_n, 0),
            "max_traction_slipping_n": round(max_traction_slipping_n, 2),
            "result_message": result_message
        }

    def _calculate_curves_common(self, calculate_shunting: bool) -> List[dict]:
        """
        Core calculation engine for performance curves.

        Generates data points for tractive effort and shunting capability curves.
        This is the heart of the performance analysis, calculating how vehicle
        performance varies with speed, slope, and gear ratio.

        Args:
            calculate_shunting (bool): If True, calculate shunting capacity (tons)
                                     If False, calculate tractive effort (N)

        Returns:
            List[dict]: Data points for plotting, each containing:
                - slope: Slope percentage
                - gear: Gear ratio used
                - speed_kmh: Speed in km/h
                - value: Tractive effort (N) or shunting capacity (tons)

        Calculation Process:
        1. Iterate over slope range (0 to max_slope)
        2. For each gear ratio, calculate speed range
        3. For each speed point, calculate available traction
        4. Calculate locomotive resistance
        5. For shunting: Calculate remaining traction for wagons
        6. For tractive effort: Return raw traction force

        Physics:
        - Engine RPM derived from vehicle speed and gear ratios
        - Torque interpolated from curve, limited by power
        - Traction capped by slipping limit
        - Resistance includes rolling, gradient, curvature, starting

        Note:
            Generates 100 speed points per slope/gear combination
            Uses numpy for efficient array operations
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")

        plot_data = []

        # Iterate over slope range in 0.5% increments
        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            for gear_ratio in self.gear_ratios:
                if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0 or gear_ratio <= 0:
                    continue

                # Calculate speed range for this gear (based on RPM limits)
                min_speed_kmh = 0.0
                max_speed_kmh = ((self.max_rpm * math.pi * self.wheel_dia_m) /
                               (gear_ratio * self.rear_axle_ratio * 60.0)) * 3.6

                # Generate 100 speed points for smooth curves
                speeds_kmh = np.linspace(min_speed_kmh, max_speed_kmh, 100)

                for speed_kmh in speeds_kmh:
                    # Convert speed to wheel RPM, then to engine RPM
                    speed_mps = speed_kmh / 3.6
                    wheel_circumference_m = math.pi * self.wheel_dia_m
                    wheel_rpm = (speed_mps / wheel_circumference_m * 60.0) if wheel_circumference_m > 0 else 0
                    engine_rpm = wheel_rpm * gear_ratio * self.rear_axle_ratio

                    # Get torque at this RPM (interpolated, power-limited)
                    torque_at_rpm = interpolate_torque(engine_rpm, self.torque_curve)
                    power_at_rpm_kw = (engine_rpm * torque_at_rpm * 2.0 * math.pi) / (60.0 * 1000.0)

                    # Limit torque by peak power
                    if power_at_rpm_kw > self.peak_power_kw and engine_rpm > 0:
                        torque_at_rpm = (self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi)

                    # Calculate available tractive force (capped by slipping)
                    max_traction_generated_n = 2.0 * (torque_at_rpm * gear_ratio * self.rear_axle_ratio) / self.wheel_dia_m
                    max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
                    actual_traction_n = min(max_traction_generated_n, max_traction_slipping_n)

                    # Calculate locomotive resistance (includes starting resistance)
                    loco_resistance = (
                        rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles)
                        + gradient_resistance(self.loco_gvw_ton, float(slope_percent))
                        + curvature_resistance(self.loco_gvw_ton, self.max_curve_deg)
                        + starting_resistance_loco(self.loco_gvw_ton)
                    )

                    y_val = 0.0
                    if calculate_shunting:
                        # Calculate shunting capacity (tons of wagons)
                        remaining_traction_n = actual_traction_n - loco_resistance
                        if remaining_traction_n > 0:
                            # Calculate resistance per ton of wagon
                            total_resistance_1ton_wagon = (
                                rolling_resistance_wagon(speed_kmh, 1)
                                + gradient_resistance(1, float(slope_percent))
                                + curvature_resistance(1, self.max_curve_deg)
                                + starting_resistance_wagon(1)
                            )
                            if total_resistance_1ton_wagon > 0:
                                y_val = remaining_traction_n / total_resistance_1ton_wagon
                    else:
                        # Return raw tractive effort (N)
                        y_val = actual_traction_n

                    plot_data.append({
                        "slope": round(slope_percent, 2),
                        "gear": gear_ratio,
                        "speed_kmh": round(speed_kmh, 2),
                        "value": round(y_val, 2)
                    })
        return plot_data

    def calculate_plot_data(self) -> dict:
        """
        Generate data for performance curves required by frontend.

        Creates the two main performance plots:
        1. Tractive Effort vs Speed (by slope and gear)
        2. Shunting Capacity vs Speed (by slope and gear)

        Returns:
            dict: Plot data for frontend consumption:
                - tractive_effort_plot: Raw tractive force curves
                - shunting_capability_plot: Wagon capacity curves

        Note:
            Both plots use the same calculation engine but different output modes
            Data format optimized for JavaScript plotting libraries
        """
        tractive_effort_data = self._calculate_curves_common(calculate_shunting=False)
        shunting_capability_data = self._calculate_curves_common(calculate_shunting=True)

        return {
            "tractive_effort_plot": tractive_effort_data,
            "shunting_capability_plot": shunting_capability_data
        }

    def calculate_speed_for_shunting_load(self) -> List[dict]:
        """
        Calculate maximum achievable speed vs slope for given shunting load.

        Generates a table showing how maximum speed decreases with increasing slope
        when pulling a specific shunting load. Useful for operating guidelines.

        Returns:
            List[dict]: Speed vs slope data, each entry containing:
                - slope: Slope percentage
                - max_speed_kmh: Maximum achievable speed at this slope

        Method:
            For each slope, finds the highest speed where traction equals resistance
            Uses maximum gear ratio (for lowest speed, highest torque)
            Considers both locomotive and wagon resistance

        Physics:
            Equilibrium: Traction Force = Total Resistance
            Total resistance = Loco resistance + Wagon resistance
            Speed found by testing 100 speed points and finding maximum where
            available traction >= required resistance

        Note:
            Uses running resistance (not starting) as per original implementation
            Assumes constant shunting load across all slopes
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")

        table_data = []

        # Calculate speed range using gear ratios
        if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0:
             raise ValueError("Axle Ratio or Wheel Diameter cannot be zero.")

        # Speed range based on RPM limits and gear ratios
        min_speed_kmh = (self.min_rpm * math.pi * self.wheel_dia_m) / (max(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6
        max_speed_kmh = (self.max_rpm * math.pi * self.wheel_dia_m) / (min(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6

        # Generate speed test points
        speeds_kmh = np.linspace(min_speed_kmh, max_speed_kmh, 100)

        # Calculate for each slope increment
        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            max_achievable_speed = 0.0

            # Test each speed to find maximum achievable
            for speed_kmh in speeds_kmh:
                # Convert speed to engine RPM (using max gear ratio)
                speed_mps = speed_kmh / 3.6
                wheel_circ = math.pi * self.wheel_dia_m
                wheel_rpm = (speed_mps / wheel_circ * 60.0) if wheel_circ > 0 else 0
                engine_rpm = wheel_rpm * max(self.gear_ratios) * self.rear_axle_ratio

                # Get torque at this RPM
                tq = interpolate_torque(engine_rpm, self.torque_curve)

                # Limit by peak power
                power_kw = (engine_rpm * tq * 2.0 * math.pi) / (60.0 * 1000.0)
                if power_kw > self.peak_power_kw and engine_rpm > 0:
                    tq = (self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi)

                # Calculate available traction (capped by slipping)
                max_traction_generated_n = 2 * (tq * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m
                max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
                actual_traction_n = min(max_traction_generated_n, max_traction_slipping_n)

                # Calculate total resistance (locomotive + wagons)
                # NOTE: Uses RUNNING resistance, not starting (as per original logic)
                loco_res = (
                    rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles)
                    + gradient_resistance(self.loco_gvw_ton, float(slope_percent))
                    + curvature_resistance(self.loco_gvw_ton, self.max_curve_deg)
                )

                wagon_res = (
                    rolling_resistance_wagon(speed_kmh, self.shunting_load_t)
                    + gradient_resistance(self.shunting_load_t, float(slope_percent))
                    + curvature_resistance(self.shunting_load_t, self.max_curve_deg)
                )

                total_resistance_n = loco_res + wagon_res

                # Check if traction can overcome resistance
                if actual_traction_n >= total_resistance_n:
                    max_achievable_speed = float(speed_kmh)

            table_data.append({
                "slope": round(slope_percent, 2),
                "max_speed_kmh": round(max_achievable_speed, 2)
            })
        return table_data

    def create_report_docx(self):
        """
        Generate comprehensive DOCX report for vehicle performance analysis.

        Creates a professional report containing all input parameters,
        calculation results, and performance tables. Report includes:
        - Input parameters table
        - Torque curve data
        - Traction analysis results
        - Speed vs slope performance table

        Returns:
            io.BytesIO: In-memory DOCX file stream

        Report Structure:
        1. Title and timestamp
        2. Input parameters table
        3. Torque curve table
        4. Traction analysis results
        5. Speed vs slope performance table

        Dependencies:
            - python-docx library for document creation
            - All calculation methods for data generation

        Note:
            Report is generated in memory and returned as byte stream
            Suitable for web API responses or file downloads
        """
        if docx is None:
            raise ImportError("python-docx library is required to generate .docx files.")

        doc = docx.Document()
        doc.add_heading("Vehicle Performance Calculator Report", level=1)
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading("Input Parameters", level=2)

        # Use the original validated inputs for the report
        inputs = self.inputs

        # Create input parameters table
        inputs_table_data = {
            "Loco GVW (kg)": inputs.get('loco_gvw_kg'),
            "Max Speed (km/h)": inputs.get('max_speed_kmh'),
            "Number of Axles": inputs.get('num_axles'),
            "Rear Axle Ratio": inputs.get('rear_axle_ratio'),
            "Gear Ratios": ", ".join(map(str, inputs.get('gear_ratios', []))),
            "Shunting Load (tons)": inputs.get('shunting_load_t'),
            "Peak Power (kW)": inputs.get('peak_power_kw'),
            "Coeff. of Friction (mu)": inputs.get('friction_mu'),
            "Wheel Dia (m)": inputs.get('wheel_dia_m'),
            "Min RPM": inputs.get('min_rpm'),
            "Max RPM": inputs.get('max_rpm'),
            "Max Curve": f"{inputs.get('max_curve')} ({inputs.get('curve_unit')})",
            "Max Slope": f"{inputs.get('max_slope')} ({inputs.get('slope_unit')})",
        }

        # Add inputs table to document
        t_inputs = doc.add_table(rows=1, cols=2)
        t_inputs.style = 'Table Grid'
        t_inputs.rows[0].cells[0].text = "Parameter"
        t_inputs.rows[0].cells[1].text = "Value"
        for label, value in inputs_table_data.items():
            row_cells = t_inputs.add_row().cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(value)

        # Torque curve section
        doc.add_heading("Torque Curve", level=3)
        t_torque = doc.add_table(rows=1, cols=2)
        t_torque.style = 'Table Grid'
        t_torque.rows[0].cells[0].text = "RPM"
        t_torque.rows[0].cells[1].text = "Torque (Nm)"
        for rpm, torque in sorted(self.torque_curve.items()):
            row_cells = t_torque.add_row().cells
            row_cells[0].text = str(rpm)
            row_cells[1].text = str(torque)

        doc.add_heading("Calculation Results", level=2)

        # Traction analysis results
        traction_results = self.run_tractive_calculation()
        doc.add_paragraph(f"Max Traction Produced: {traction_results['max_traction_generated_n']:.2f} N")
        doc.add_paragraph(f"Max Traction (No Slip): {traction_results['max_traction_slipping_n']:.2f} N")
        p = doc.add_paragraph()
        run = p.add_run(f"Result: {traction_results['result_message']}")
        run.font.bold = True

        # Speed vs slope performance table
        doc.add_heading("Max Achievable Speed vs. Slope", level=3)
        doc.add_paragraph(f"(For Shunting Load: {self.shunting_load_t} tons)")

        table_data = self.calculate_speed_for_shunting_load()
        t_speed = doc.add_table(rows=1, cols=2)
        t_speed.style = 'Table Grid'
        t_speed.rows[0].cells[0].text = "Slope (%)"
        t_speed.rows[0].cells[1].text = "Max Achievable Speed (km/h)"
        for row in table_data:
            row_cells = t_speed.add_row().cells
            row_cells[0].text = f"{row['slope']:.2f}"
            row_cells[1].text = f"{row['max_speed_kmh']:.2f}"

        # Save to in-memory stream and return
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream


def perform_vehicle_performance_calculation(inputs: Dict[str, Any]) -> Dict[str, Any]:
    """
    Main orchestration function for vehicle performance calculations.

    This is the primary entry point for vehicle performance analysis.
    Creates a calculator instance and runs all performance calculations.

    Args:
        inputs (Dict[str, Any]): Complete input parameters including:
            - Vehicle parameters (weight, dimensions, powertrain)
            - Track parameters (curves, slopes)
            - Operating parameters (loads, speeds)

    Returns:
        Dict[str, Any]: Complete analysis results containing:
            - traction_snapshot: Traction vs slipping analysis
            - plot_data: Performance curves for plotting
            - speed_slope_table: Speed vs slope lookup table

    Process Flow:
        1. Create VehiclePerformanceCalculator instance
        2. Run traction analysis
        3. Generate plot data for curves
        4. Calculate speed vs slope table
        5. Return comprehensive results

    Note:
        This function coordinates all calculations and serves as the
        main interface between API layer and calculation engine.
    """
    calculator = VehiclePerformanceCalculator(inputs)

    results = {
        'traction_snapshot': calculator.run_tractive_calculation(),
        'plot_data': calculator.calculate_plot_data(),
        'speed_slope_table': calculator.calculate_speed_for_shunting_load()
    }

    return results