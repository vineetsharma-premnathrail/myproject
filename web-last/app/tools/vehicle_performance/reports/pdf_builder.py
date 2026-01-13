# =========================================
# Vehicle Performance Tool DOCX Report Builder
# =========================================
# Professional engineering report generation for locomotive performance analysis.
#
# Purpose: Generate comprehensive Microsoft Word documents containing detailed
# locomotive performance analysis results, specifications, and engineering data
# for railway vehicle procurement, design review, and performance evaluation.
#
# Engineering Context:
# - Creates formatted DOCX reports for locomotive performance specifications
# - Includes complete input parameters, traction analysis, and performance curves
# - Critical for engineering documentation, procurement decisions, and compliance
# - Supports railway engineering standards and professional report formatting
#
# Report Structure:
# - Executive summary with locomotive specifications
# - Detailed traction performance analysis by gear
# - Speed vs slope performance characteristics
# - Shunting capability analysis for yard operations
#
# Used by: vehicle_performance/api.py (download-report endpoint)
# Dependencies: python-docx library for document generation
# Output: In-memory DOCX stream for web download

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List

def create_vehicle_performance_docx_report(inputs: Dict[str, Any], results: Dict[str, Any]) -> io.BytesIO:
    """
    Generate comprehensive locomotive performance analysis report as DOCX document.

    Creates a professional Microsoft Word document containing complete engineering
    analysis of locomotive performance characteristics, including traction curves,
    speed capabilities, and operational performance data for railway applications.

    Engineering Documentation:
    - Structured report format compliant with railway engineering standards
    - Includes timestamp for version control and audit trails
    - Comprehensive input specifications and analysis results
    - Suitable for procurement, design review, and regulatory submissions

    Parameters:
        inputs (dict): Validated locomotive and transmission specifications
                      - loco_gvw_kg: Gross vehicle weight (kg)
                      - peak_power_kw: Maximum engine power (kW)
                      - wheel_dia_m: Drive wheel diameter (m)
                      - num_axles: Number of powered axles
                      - rear_axle_ratio: Final drive gear ratio
                      - gear_ratios: List of transmission gear ratios
                      - friction_mu: Wheel-rail friction coefficient
                      - max_speed_kmh: Maximum design speed (km/h)
                      - max_slope: Maximum operating gradient (%)
                      - max_curve: Maximum curve angle (degrees)
                      - shunting_load_t: Maximum shunting load (tonnes)

        results (dict): Complete performance analysis results from calculator
                       - traction_snapshot: Performance summary by gear
                       - plot_data: Graph data for performance curves
                       - speed_slope_table: Speed vs gradient analysis

    Returns:
        io.BytesIO: In-memory DOCX file stream ready for download
                   - Contains formatted engineering report
                   - Compatible with Microsoft Word and other DOCX readers
                   - Includes all performance analysis data and specifications

    Report Sections:

    1. Input Parameters:
       - Complete locomotive specifications (weight, power, dimensions)
       - Transmission configuration (gear ratios, axle ratios)
       - Operating constraints (speed, gradient, curve limits)
       - Performance requirements (shunting capacity)

    2. Traction Performance Snapshot:
       - Maximum speeds for each gear on level track
       - Performance on design slope and curve conditions
       - Tractive effort capabilities by gear ratio
       - Critical for route performance analysis

    3. Speed vs Slope Analysis:
       - Performance degradation with increasing gradients
       - Speed capabilities across different gear ratios
       - Essential for route planning and timetable development
       - Shows first few and last entries for comprehensive analysis

    4. Shunting Performance:
       - Low-speed performance for yard operations
       - Maximum shunting loads and corresponding speeds
       - Important for locomotive utilization and capacity planning

    Technical Implementation:
    - Uses python-docx library for professional document formatting
    - Creates tables with proper column widths and headers
    - Includes timestamp for report versioning
    - Handles missing data gracefully with appropriate messages

    Error Handling:
    - Catches document creation errors and provides detailed messages
    - Validates data availability before creating report sections
    - Returns RuntimeError for any document generation failures

    Example:
        >>> inputs = {
        ...     'loco_gvw_kg': 120000,
        ...     'peak_power_kw': 2000,
        ...     'wheel_dia_m': 1.25,
        ...     'num_axles': 4,
        ...     'rear_axle_ratio': 3.5,
        ...     'gear_ratios': [8.5, 6.2, 4.5, 3.2, 2.1],
        ...     'friction_mu': 0.35,
        ...     'max_speed_kmh': 100,
        ...     'max_slope': 2.0,
        ...     'max_curve': 6.0,
        ...     'shunting_load_t': 2000
        ... }
        >>> results = calculator.get_analysis_results()
        >>> docx_stream = create_vehicle_performance_docx_report(inputs, results)
        >>> # Returns BytesIO stream containing formatted DOCX report

    Dependencies:
        - python-docx: Required for DOCX document creation
        - datetime: For report timestamp generation
        - io: For in-memory file stream handling

    Note:
        Report generation is synchronous and may take several seconds
        for complex performance analyses with multiple gear ratios.
        All calculations must be completed before report generation begins.
    """
    # Initialize new Word document for professional engineering report
    doc = Document()

    # Add main report title and generation timestamp
    doc.add_heading('Vehicle Performance Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Complete locomotive and transmission specifications
    doc.add_heading('1. Input Parameters', level=1)
    p = doc.add_paragraph()

    # Add locomotive weight and power specifications
    p.add_run(f"  • Locomotive GVW: {inputs['loco_gvw_kg']:.0f} kg\n")
    p.add_run(f"  • Peak Power: {inputs['peak_power_kw']:.0f} kW ({inputs['peak_power_kw'] * 1.341:.0f} HP)\n")

    # Add mechanical specifications (wheel diameter, axles, transmission)
    p.add_run(f"  • Wheel Diameter: {inputs['wheel_dia_m']:.2f} m\n")
    p.add_run(f"  • Number of Axles: {inputs['num_axles']}\n")
    p.add_run(f"  • Rear Axle Ratio: {inputs['rear_axle_ratio']:.2f}\n")
    p.add_run(f"  • Gear Ratios: {', '.join([f'{r:.2f}' for r in inputs.get('gear_ratios', [])])}\n")

    # Add operational parameters (friction, speed limits, route constraints)
    p.add_run(f"  • Friction Coefficient: {inputs.get('friction_mu', 0.3):.2f}\n")
    p.add_run(f"  • Maximum Speed: {inputs['max_speed_kmh']:.0f} km/h\n")
    p.add_run(f"  • Maximum Slope: {inputs.get('max_slope', 0):.1f}%\n")
    p.add_run(f"  • Maximum Curve: {inputs.get('max_curve', 0):.0f} degrees\n")

    # Add shunting capacity if specified
    if inputs.get('shunting_load_t'):
        p.add_run(f"  • Shunting Load: {inputs['shunting_load_t']:.1f} tons\n")

    # Section 2: Traction performance summary table
    doc.add_heading('2. Traction Performance Snapshot', level=1)
    traction_snapshot = results.get('traction_snapshot', {})

    if traction_snapshot:
        # Create performance table with proper formatting
        table = doc.add_table(rows=1, cols=4)
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1.5)
        table.columns[3].width = Inches(1.5)

        # Add table headers for performance metrics
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Gear Ratio'
        hdr_cells[1].text = 'Max Speed (Level) km/h'
        hdr_cells[2].text = 'Max Speed (Slope) km/h'
        hdr_cells[3].text = 'Max Speed (Curve) km/h'

        # Add performance data for each gear ratio
        for gear_key, gear_data in traction_snapshot.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f"{gear_data['gear_ratio']:.2f}"
            row_cells[1].text = f"{gear_data['max_speed_level_kmh']:.1f}"
            row_cells[2].text = f"{gear_data['max_speed_slope_kmh']:.1f}"
            row_cells[3].text = f"{gear_data['max_speed_curve_kmh']:.1f}"
    else:
        # Handle case where traction data is not available
        doc.add_paragraph("No traction data available.")

    # Section 3: Speed vs slope performance analysis
    doc.add_heading('3. Speed vs Slope Analysis', level=1)
    speed_slope_table = results.get('speed_slope_table', [])

    if speed_slope_table:
        # Create table for gradient performance analysis
        table = doc.add_table(rows=1, cols=len(speed_slope_table[0]) if speed_slope_table else 1)
        table.columns[0].width = Inches(1.0)

        # Add table headers with gear ratio labels
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Slope (%)'
        col_idx = 1
        for key in speed_slope_table[0].keys():
            if key != 'slope_percent':
                gear_ratio = key.replace('gear_', '').replace('_speed_kmh', '')
                hdr_cells[col_idx].text = f'Gear {gear_ratio} (km/h)'
                col_idx += 1

        # Add performance data (show representative samples for readability)
        # Display every other row for first 5 entries, plus the last entry
        display_rows = speed_slope_table[::2][:5] + speed_slope_table[-1:] if len(speed_slope_table) > 6 else speed_slope_table
        for row_data in display_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{row_data['slope_percent']:.1f}"
            col_idx = 1
            for key, value in row_data.items():
                if key != 'slope_percent':
                    row_cells[col_idx].text = f"{value:.1f}"
                    col_idx += 1
    else:
        # Handle case where speed-slope data is not available
        doc.add_paragraph("No speed vs slope data available.")

    # Section 4: Shunting performance analysis
    doc.add_heading('4. Shunting Performance', level=1)
    plot_data = results.get('plot_data', {})
    shunting_data = plot_data.get('shunting_capability_plot', {})

    if shunting_data and shunting_data.get('loads'):
        # Add shunting capacity summary
        doc.add_paragraph(f"Shunting Load: {shunting_data['loads'][0]:.1f} tons")
        doc.add_paragraph(f"Maximum Speed: {shunting_data['speeds'][0]:.1f} km/h")
    else:
        # Handle case where shunting data is not available
        doc.add_paragraph("No shunting performance data available.")

    try:
        # Save completed document to in-memory stream for web download
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        # Handle document creation errors with detailed error messages
        raise RuntimeError(f"An error occurred while creating the report: {e}")