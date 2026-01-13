"""
Hydraulic Tool DOCX Report Builder Module
========================================
Purpose:
    Generates professional Microsoft Word documents for hydraulic calculations.
    Creates downloadable engineering reports with formatted tables and professional layout.
Layer:
    Backend / Tools / Hydraulic / Reports
Technology:
    - python-docx: Microsoft Word document generation library
    - Professional document formatting with tables, headings, and styling
    - In-memory document creation and streaming
Process:
    1. Create new Word document
    2. Add title page with generation timestamp
    3. Create input parameter tables
    4. Add step-by-step calculation sections
    5. Include results tables and analysis
    6. Return document as downloadable stream
Dependencies:
    - python-docx library (optional import with graceful fallback)
    - Valid calculation results and input data
Report Types:
    - Displacement calculation reports (calc_cc mode)
    - Speed calculation reports (calc_speed mode)
Features:
    - Professional engineering document formatting
    - Structured tables for input parameters
    - Step-by-step calculation breakdowns
    - Color-coded results and compliance indicators
    - Automatic timestamp and metadata
Error Handling:
    - Graceful handling of missing python-docx library
    - Clear error messages for installation instructions
    - Validation of input data before document creation
"""

import io
from datetime import datetime
from typing import Dict, Any

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

def create_hydraulic_docx_report(inputs: Dict[str, Any], results: Dict[str, Any], inputs_raw: Dict[str, Any]) -> io.BytesIO:
    """
    Create a professional DOCX report for hydraulic calculations.

    Generates a comprehensive Microsoft Word document containing all input parameters,
    calculations, and results in a professionally formatted engineering report.

    Args:
        inputs (Dict[str, Any]): Processed input parameters
        results (Dict[str, Any]): Detailed calculation results
        inputs_raw (Dict[str, Any]): Original raw input data for display

    Returns:
        io.BytesIO: In-memory buffer containing the DOCX file
            Ready for HTTP streaming/download

    Raises:
        ImportError: If python-docx library is not installed
        Exception: For document creation or data processing errors

    Document Structure:
        - Title page with report type and generation timestamp
        - Input parameters section with organized tables
        - Step-by-step calculations with formulas and intermediate results
        - Results tables with motor/pump specifications
        - Performance analysis and recommendations
        - Professional formatting with headers, tables, and styling

    Dependencies:
        - python-docx: Required for Word document generation
        - Install with: pip install python-docx

    File Format:
        - Microsoft Word DOCX format
        - Compatible with Word 2007 and later
        - Professional engineering document standards
    """
    if docx is None:
        raise ImportError("python-docx library is required to generate .docx files. Install with: pip install python-docx")

    # Create new Word document
    doc = docx.Document()

    # Route to appropriate report generation function based on calculation mode
    if inputs['calc_mode'] == "calc_cc":
        _create_displacement_docx(doc, inputs, results, inputs_raw)
    else:
        _create_speed_docx(doc, inputs, results, inputs_raw)

    # Save document to in-memory buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Reset stream position for reading
    return file_stream

def _create_displacement_docx(doc, inputs, results, raw):
    """
    Create DOCX report for displacement calculation mode.

    Generates a detailed report for motor displacement and pump sizing calculations,
    including all input parameters, calculations, and specifications.

    Args:
        doc: python-docx Document object
        inputs: Processed input parameters
        results: Calculation results
        raw: Original raw input data

    Document Sections:
        1. Title and timestamp
        2. Vehicle input parameters table
        3. Hydraulic system input parameters table
        4. Step-by-step calculation breakdown
        5. Motor specifications and requirements
        6. Pump specifications for each gear ratio
        7. Performance analysis and recommendations
    """
    doc.add_heading("Pump & Motor (cc) Calculation Report", level=0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Input tables
    doc.add_heading("Input Parameters", level=2)
    vehicle_data = {
        "Vehicle Weight (t)": raw.get('weight'), "Number of axles": raw.get('axles'),
        "Target Speed (km/h)": raw.get('speed'), "Slope (%)": raw.get('slope_percent'),
        "Curve (deg)": raw.get('curve_degree'), "Wheel Dia (mm)": raw.get('wheel_diameter'),
        "PTO Gear box Ratio": raw.get('pto_gear_ratio'), "Axle Gear box Ratio": raw.get('axle_gear_box_ratio'),
        "Engine Gear Box Ratios": raw.get('engine_gear_ratio'), "max Vehicle RPM": raw.get('max_vehicle_rpm'),
    }
    hydraulic_data = {
        "Hydraulic Motor (Total)": raw.get('num_motors'), "Hydraulic Motor / axle": raw.get('per_axle_motor'),
        "Pressure (bar)": raw.get('pressure'), "Motor Mech Eff (%)": raw.get('mech_eff_motor'),
        "Motor Vol Eff (%)": raw.get('vol_eff_motor'), "Pump Vol Eff (%)": raw.get('vol_eff_pump'),
    }
    _populate_input_table(doc, "Vehicle Inputs", vehicle_data)
    _populate_input_table(doc, "Hydraulic Inputs", hydraulic_data)

    doc.add_page_break()
    doc.add_heading("Calculation Results", level=2)

    # Summary results
    p = doc.add_paragraph("Motor Displacement: ")
    run = p.add_run(f"{results.get('motor_displacement_cc'):.2f} cc/rev")
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0D47A1")

    p = doc.add_paragraph("Motor Flow Rate: ")
    run = p.add_run(f"{results.get('per_motor_flow_rate_lpm'):.2f} LPM")
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0D47A1")

    # Pump results table
    doc.add_heading("Required Pump Displacement (Per Engine Gear)", level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Engine Gear Ratio'
    hdr_cells[1].text = 'max Vehicle RPM'
    hdr_cells[2].text = 'Calculated Pump RPM'
    hdr_cells[3].text = 'Required Pump Disp. (cc/rev)'

    pump_results_list = results.get('pump_results', [])
    for res in pump_results_list:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{res.get('engine_gear_ratio', 0):.2f}"
        row_cells[1].text = f"{res.get('max_vehicle_rpm_input', 0):.0f}"
        row_cells[2].text = f"{res.get('pump_rpm', 0):.2f}"
        run = row_cells[3].paragraphs[0].add_run(f"{res.get('pump_disp_cc', 0):.2f}")
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("0D47A1")

def _create_speed_docx(doc, inputs, results, raw):
    """Create DOCX for speed calculation mode"""
    doc.add_heading("Achievable Speed Calculation Report", level=0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Input tables
    doc.add_heading("Input Parameters", level=2)
    vehicle_data = {
        "Wheel Dia (mm)": raw.get('wheel_diameter'), "PTO Gear box Ratio": raw.get('pto_gear_ratio'),
        "Axle Gear box Ratio": raw.get('axle_gear_box_ratio'), "Engine Gear Box Ratios": raw.get('engine_gear_ratio'),
        "max Vehicle RPM": raw.get('max_vehicle_rpm'),
    }
    hydraulic_data = {
        "Hydraulic Motor (Total)": raw.get('num_motors'), "Hydraulic Motor / axle": raw.get('per_axle_motor'),
        "Pump Displacement (cc)": raw.get('pump_disp_in'), "Pump Vol Eff (%)": raw.get('vol_eff_pump'),
        "max Limit of Pump RPM": raw.get('max_pump_rpm'), "Motor Displacement (cc)": raw.get('motor_disp_in'),
        "Motor Vol Eff (%)": raw.get('vol_eff_motor'), "max Motor RPM": raw.get('max_motor_rpm'),
    }
    _populate_input_table(doc, "Vehicle Inputs", vehicle_data)
    _populate_input_table(doc, "Hydraulic Inputs", hydraulic_data)

    # Results table
    doc.add_heading("Calculation Results (Summary)", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Engine Gear Ratio'
    hdr_cells[1].text = 'max Vehicle RPM'
    hdr_cells[2].text = 'Calc. Pump RPM'
    hdr_cells[3].text = 'Calc. Pump Flow (LPM)'
    hdr_cells[4].text = 'Calc. Motor Speed (RPM)'
    hdr_cells[5].text = 'Calc. Axle/Wheel Speed (RPM)'
    hdr_cells[6].text = 'Achievable Speed (km/h)'

    speed_results_list = results.get('speed_results_list', [])
    for res in speed_results_list:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{res.get('engine_gear_ratio', 0):.2f}"
        row_cells[1].text = f"{res.get('max_vehicle_rpm_input', 0):.0f}"
        calc_pump_rpm = res.get('pump_rpm', 0)
        run_pump_rpm = row_cells[2].paragraphs[0].add_run(f"{calc_pump_rpm:.2f}")
        if calc_pump_rpm > inputs.get('max_pump_rpm'):
            run_pump_rpm.font.color.rgb = RGBColor.from_string("FF0000")
            run_pump_rpm.font.bold = True
        row_cells[3].text = f"{res.get('pump_flow_lpm', 0):.2f}"
        calc_motor_rpm = res.get('motor_speed_rpm', 0)
        run_motor_rpm = row_cells[4].paragraphs[0].add_run(f"{calc_motor_rpm:.2f}")
        if calc_motor_rpm > inputs.get('max_motor_rpm'):
            run_motor_rpm.font.color.rgb = RGBColor.from_string("FF0000")
            run_motor_rpm.font.bold = True
        row_cells[5].text = f"{res.get('axle_shaft_rpm', 0):.2f}"
        run = row_cells[6].paragraphs[0].add_run(f"{res.get('achievable_speed_kph', 0):.2f}")
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("0D47A1")

    # Warnings
    if results.get('warnings'):
        p_note = doc.add_paragraph("Note: Values in ")
        run_red = p_note.add_run("red")
        run_red.font.color.rgb = RGBColor.from_string("FF0000")
        run_red.font.bold = True
        p_note.add_run(f" exceed the specified limits (Max Motor: {inputs.get('max_motor_rpm'):.0f} RPM, Max Pump: {inputs.get('max_pump_rpm'):.0f} RPM).")

def _populate_input_table(doc, title, data_dict):
    """Helper to populate input parameter tables"""
    if docx is None:
        return
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(data_dict.items()):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(value)
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    return table