# =========================================
# Qmax Tool DOCX Report Builder Module
# =========================================
# Professional document generation for Qmax engineering reports.
#
# Purpose: Create formatted Microsoft Word documents containing complete
# Qmax calculation reports with proper engineering documentation standards,
# suitable for regulatory submissions, design reviews, and professional records.
#
# Engineering Context:
# - Generates professional engineering documentation
# - Ensures consistent formatting across all reports
# - Supports railway industry standards for calculation documentation
# - Provides timestamped, auditable calculation records
#
# Used by: api.py (/download_qmax_report endpoint)
# Dependencies: python-docx (document generation), constants.py (CONSTANT_C)

import io
from datetime import datetime
from typing import Dict, Any

try:
    import docx
    from docx.shared import Pt, RGBColor
except ImportError:
    docx = None

from ..constants import CONSTANT_C

def create_qmax_docx_report(results: Dict[str, Any], inputs_raw: Dict[str, Any]) -> io.BytesIO:
    """
    Create a professional DOCX report for Qmax calculations.

    Generates a formatted Microsoft Word document containing the complete
    engineering calculation report with input parameters, step-by-step
    methodology, and final results. Designed for railway engineering
    documentation and regulatory compliance.

    Document Structure:
    1. Title and timestamp header
    2. Input parameters section with engineering context
    3. Detailed calculation methodology with formulas
    4. Final results with appropriate units and formatting

    Args:
        results (Dict[str, Any]): Complete calculation results from service layer
                                 containing all inputs and computed Qmax values

        inputs_raw (Dict[str, Any]): Original user inputs for display in report
                                    - Preserves user selections and custom values
                                    - Maintains context for engineering records

    Returns:
        io.BytesIO: In-memory byte stream containing the complete DOCX file
                   - Ready for HTTP response streaming
                   - Compatible with FastAPI StreamingResponse
                   - No temporary files created on disk

    Document Sections:

    HEADER:
    - Report title: "Qmax Calculation Report"
    - Generation timestamp for version control and audit trails

    INPUT PARAMETERS:
    - Wheel diameter limit with units (mm)
    - Material strength selection and actual value used (N/mm²)
    - Safety factor for dynamic loading conditions

    CALCULATION METHODOLOGY:
    - Engineering formula with empirical constants
    - Parameter substitution with proper units
    - Step-by-step mathematical derivation
    - Intermediate calculation values for verification

    FINAL RESULTS:
    - Qmax in kiloNewtons (primary engineering unit)
    - Qmax in metric tonnes (practical reference)
    - Highlighted formatting for easy identification

    Formatting Features:
    - Professional document styling with headings
    - Bold formatting for important results
    - Blue color highlighting for final Qmax values
    - Structured layout for engineering readability

    Raises:
        ImportError: If python-docx library is not installed
                    - Graceful degradation with clear error message
                    - Allows application to continue without DOCX functionality

    Example:
        >>> results = {'d': 750.0, 'sigma_b': 880.0, 'v_head': 1.2,
        ...           'qmax_kn': 123.45, 'qmax_tonnes': 12.58}
        >>> inputs_raw = {'d': '750', 'sigma_b_selection': '880 N/mm²', 'v_head': '1.2'}
        >>> docx_stream = create_qmax_docx_report(results, inputs_raw)
        >>> # Returns BytesIO stream ready for download

    Technical Notes:
    - Uses in-memory streams to avoid disk I/O
    - Compatible with cloud deployment environments
    - Thread-safe document generation
    - Minimal memory footprint for large-scale operations

    Engineering Standards Compliance:
    - Follows railway engineering documentation practices
    - Includes all necessary calculation parameters
    - Provides traceable calculation methodology
    - Suitable for regulatory review and approval processes
    """

    # Verify python-docx availability for document generation
    if docx is None:
        raise ImportError("python-docx library is required to generate .docx files.")

    # Create new Word document with professional formatting
    doc = docx.Document()

    # Add report header with title and timestamp
    doc.add_heading('Qmax Calculation Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Input Parameters Section
    doc.add_heading('1. Input Parameters', level=1)
    p = doc.add_paragraph()
    p.add_run(f"  • Worn rail diameter limit (d): {inputs_raw.get('d')} mm\n")
    p.add_run(f"  • Material Strength (σB): {inputs_raw.get('sigma_b_selection')}\n")
    p.add_run(f"    (Value Used: {results['sigma_b']} N/mm²)\n")
    p.add_run(f"  • Safety Factor (v_head): {inputs_raw.get('v_head')}\n")

    # Calculation Methodology Section
    doc.add_heading('2. Calculation', level=1)

    # Calculate intermediate values for detailed reporting
    sigma_v_head_squared = (results['sigma_b'] / results['v_head']) ** 2
    d_half = results['d'] / 2

    # Format detailed calculation steps with engineering precision
    detailed_steps_text = (
        f"1. Formula:\n"
        f"   Qmax = C × (d / 2) × (σB / v_head)²\n"
        f"   Where C = {CONSTANT_C}\n\n"
        f"2. Substitute Values:\n"
        f"   d = {results['d']} mm\n"
        f"   σB = {results['sigma_b']} N/mm²\n"
        f"   v_head = {results['v_head']}\n\n"
        f"3. Step-by-Step Calculation:\n"
        f"   a) (σB / v_head)² = ({results['sigma_b']} / {results['v_head']})² = {sigma_v_head_squared:.3f}\n\n"
        f"   b) Qmax = {CONSTANT_C} × ({results['d']} / 2) × {sigma_v_head_squared:.3f}\n"
        f"   c) Qmax = {CONSTANT_C} × {d_half:.1f} × {sigma_v_head_squared:.3f}\n"
    )
    doc.add_paragraph(detailed_steps_text)

    # Final Results Section with professional formatting
    doc.add_heading('3. Final Result', level=1)
    p_res = doc.add_paragraph()

    # Highlight primary result in kN with bold blue formatting
    run_kn = p_res.add_run(f"Qmax = {results['qmax_kn']:.4f} kN\n")
    run_kn.font.bold = True
    run_kn.font.color.rgb = RGBColor.from_string("0D47A1")  # Professional blue color

    # Add secondary result in tonnes
    p_res.add_run(f"Qmax = {results['qmax_tonnes']:.4f} tonnes")

    # Save document to in-memory stream for efficient delivery
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Reset stream position for reading

    return file_stream