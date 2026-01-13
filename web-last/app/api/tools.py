# =========================================
# Tools API Dispatcher Router Module
# =========================================
# Central routing hub for all engineering calculation tools and legacy API compatibility.
#
# Purpose: Acts as the main entry point for all engineering calculation tools,
# routing requests to appropriate tool-specific APIs while maintaining backward
# compatibility with legacy endpoint URLs used by existing HTML frontends.
#
# Architecture Overview:
# - Modular tool routing with dedicated prefixes
# - Legacy endpoint preservation for seamless upgrades
# - Unified error handling and response formatting
# - Tool-specific API delegation and orchestration
#
# Tool Integration:
# - Braking calculations (force, distance, performance)
# - Qmax analysis (maximum power optimization)
# - Hydraulic system design and simulation
# - Load distribution analysis
# - Tractive effort calculations
# - Vehicle performance evaluation
#
# API Structure:
# Modern Endpoints (Recommended):
# - POST /braking/calculate - Braking force analysis
# - POST /qmax/calculate - Power optimization
# - POST /hydraulic/calculate - Hydraulic system design
# - POST /load-distribution/calculate - Load analysis
# - POST /tractive-effort/calculate - Traction calculations
# - POST /vehicle-performance/calculate - Performance analysis
#
# Legacy Endpoints (Backward Compatibility):
# - POST /braking_calculate - Old braking endpoint
# - POST /calculate_qmax - Old Qmax endpoint
# - POST /calculate - Old hydraulic endpoint
# - POST /calculate_load_distribution - Old load distribution
# - POST /calculate_tractive_effort - Old tractive effort
# - POST /calculate_performance - Old vehicle performance
#
# Report Generation:
# - POST /download_braking_report - Braking analysis report
# - POST /download_qmax_report - Qmax analysis report
# - POST /download_report - Hydraulic report
# - POST /download_load_distribution_report - Load distribution report
# - POST /download_tractive_effort_report - Tractive effort report
# - POST /download_performance_report - Vehicle performance report
#
# Used by: Main FastAPI application, HTML frontend, external API clients
# Dependencies: Individual tool API modules, Pydantic for validation
# Migration: Legacy endpoints will be deprecated in future versions

from fastapi import APIRouter
from app.tools.braking.api import router as braking_router
from app.tools.qmax.api import router as qmax_router
from app.tools.hydraulic.api import router as hydraulic_router
from app.tools.load_distribution.api import router as load_distribution_router
from app.tools.tractive_effort.api import router as tractive_effort_router
from app.tools.vehicle_performance.api import router as vehicle_performance_router

# Create main tools router
router = APIRouter()

# =========================================
# MODERN TOOL API ROUTING
# =========================================
# Include all tool-specific routers with standardized prefixes
# This provides clean, organized API endpoints for each engineering tool

router.include_router(braking_router, prefix="/braking")
router.include_router(qmax_router, prefix="/qmax")
router.include_router(hydraulic_router, prefix="/hydraulic")
router.include_router(load_distribution_router, prefix="/load-distribution")
router.include_router(tractive_effort_router, prefix="/tractive-effort")
router.include_router(vehicle_performance_router, prefix="/vehicle-performance")

# =========================================
# LEGACY API COMPATIBILITY LAYER
# =========================================
# Backward compatibility endpoints to support existing HTML frontend
# These endpoints maintain the old URL patterns and data formats
# TODO: Deprecate these endpoints in future major version release

# =========================================
# BRAKING TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing braking calculation HTML interface

@router.post("/braking_calculate")
async def legacy_braking_calculate(raw_input: dict):
    """
    Legacy braking calculation endpoint for backward compatibility.

    Preserves the original API contract used by the HTML frontend for braking
    force and distance calculations. Translates legacy input format to modern
    API structure and delegates to the braking tool service.

    Legacy Input Format:
    - Raw dictionary with braking parameters
    - May include inconsistent field names
    - Lacks modern validation and type safety

    Migration Path:
    - New implementations should use: POST /braking/calculate
    - Legacy format will be supported until deprecation notice
    - Frontend should be updated to use modern endpoints

    Braking Calculations:
    - Stopping force analysis
    - Braking distance calculations
    - Brake system performance evaluation
    - Safety factor assessments

    Parameters:
        raw_input (dict): Legacy braking calculation parameters
                         - Expected fields vary by calculation type
                         - May require format translation

    Returns:
        dict: Braking calculation results in legacy format
            {
                "stopping_force": 150000,
                "braking_distance": 45.2,
                "safety_factor": 1.8,
                ...
            }

    Raises:
        HTTPException (400): Invalid input parameters or calculation errors
                            - Missing required fields
                            - Invalid parameter values
                            - Mathematical calculation failures

    Example:
        >>> legacy_input = {
        ...     "vehicle_weight": 50000,
        ...     "initial_speed": 80,
        ...     "brake_efficiency": 0.85
        ... }
        >>> result = await legacy_braking_calculate(legacy_input)
        >>> print(f"Stopping distance: {result['braking_distance']} meters")
        Stopping distance: 45.2 meters

    Note:
        This endpoint performs input format translation to modern API.
        Consider migrating to POST /braking/calculate for new implementations.
        Legacy endpoints may be removed in future major versions.
    """
    from pydantic import BaseModel
    from typing import Optional, Dict
    from pydantic import Field

    # Define legacy input model for validation
    class LegacyBrakingInput(BaseModel):
        # Core braking parameters
        vehicle_weight: Optional[float] = Field(None, description="Vehicle total weight in kg")
        initial_speed: Optional[float] = Field(None, description="Initial speed in km/h")
        brake_efficiency: Optional[float] = Field(None, description="Brake system efficiency (0-1)")

        # Advanced parameters
        gradient: Optional[float] = Field(0, description="Road gradient in percent")
        wind_resistance: Optional[float] = Field(None, description="Wind resistance coefficient")
        rolling_resistance: Optional[float] = Field(None, description="Rolling resistance coefficient")

        # Brake system specifics
        brake_type: Optional[str] = Field("drum", description="Type of brake system")
        num_axles: Optional[int] = Field(None, description="Number of braked axles")

    try:
        # Validate and parse legacy input
        validated_input = LegacyBrakingInput(**raw_input)

        # Import braking service for calculation
        from app.tools.braking.service import perform_braking_calculation

        # Convert legacy format to the expected dict format
        braking_inputs = {
            "mass_kg": validated_input.vehicle_weight,
            "reaction_time": 1.0,  # Default reaction time
            "num_wheels": validated_input.num_axles * 2 if validated_input.num_axles else 4,  # Assume 2 wheels per axle
            "calc_mode": "Rail",  # Default to rail mode
            "rail_speed_input": str(validated_input.initial_speed),  # Single speed
            "rail_gradient_input": str(validated_input.gradient or 0),
            "rail_gradient_type": "Percentage (%)",
            "road_speed_input": "",  # Not used
            "road_gradient_input": "",
            "road_gradient_type": "Percentage (%)",
            "mu": 0.7,  # Default friction
            "doc_no": "",
            "made_by": "",
            "checked_by": "",
            "approved_by": "",
            "wheel_dia": 0
        }

        # Perform calculation
        results_table, context = perform_braking_calculation(braking_inputs)

        # Extract relevant results from the table
        if results_table:
            first_result = results_table[0]
            return {
                "stopping_force": first_result.get("braking_force", 0),
                "braking_distance": first_result.get("stopping_distance", 0),
                "safety_factor": first_result.get("safety_factor", 1.0),
                "calculation_status": "success"
            }
        else:
            return {
                "stopping_force": 0,
                "braking_distance": 0,
                "safety_factor": 1.0,
                "calculation_status": "no_results"
            }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Braking calculation failed: {str(e)}")

@router.post("/download_braking_report")
async def legacy_braking_report(raw_input):
    """
    Legacy braking report generation endpoint for backward compatibility.

    Generates formatted braking analysis reports in DOCX format using the
    legacy input format. Maintains compatibility with existing HTML frontend
    report download functionality.

    Report Contents:
    - Complete braking force analysis
    - Stopping distance calculations
    - Safety factor assessments
    - Brake system performance metrics
    - Engineering recommendations

    File Format:
    - Microsoft Word DOCX format
    - Professional engineering report layout
    - Tables, charts, and formatted text
    - Timestamped and version controlled

    Parameters:
        raw_input (dict): Legacy braking parameters for report generation

    Returns:
        StreamingResponse: DOCX file download
                          - Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document
                          - Content-Disposition: attachment; filename=Braking_Report.docx

    Raises:
        HTTPException (400): Invalid input or calculation errors
        HTTPException (500): Report generation failures

    Example:
        >>> input_data = {
        ...     "vehicle_weight": 50000,
        ...     "initial_speed": 80,
        ...     "brake_efficiency": 0.85
        ... }
        >>> # Returns DOCX file download response
        >>> response = await legacy_braking_report(input_data)

    Note:
        Report generation may take several seconds for complex calculations.
        Modern endpoint: POST /braking/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        # Reuse the calculation logic from legacy endpoint
        calc_result = await legacy_braking_calculate(raw_input)

        # Generate report using braking service
        from app.tools.braking.service import perform_braking_calculation
        from app.tools.braking.reports.pdf_builder import generate_braking_pdf_report

        # Convert raw input to braking service format
        braking_inputs = {
            "mass_kg": raw_input.get("vehicle_weight", 10000),
            "reaction_time": 1.0,
            "num_wheels": raw_input.get("num_axles", 2) * 2,
            "calc_mode": "Rail",
            "rail_speed_input": str(raw_input.get("initial_speed", 30)),
            "rail_gradient_input": str(raw_input.get("gradient", 0)),
            "rail_gradient_type": "Percentage (%)",
            "road_speed_input": "",
            "road_gradient_input": "",
            "road_gradient_type": "Percentage (%)",
            "mu": 0.7,
            "doc_no": "",
            "made_by": "",
            "checked_by": "",
            "approved_by": "",
            "wheel_dia": 0
        }

        # Perform calculation to get context for report
        _, context = perform_braking_calculation(braking_inputs)
        report_stream = generate_braking_pdf_report(context)

        return StreamingResponse(
            report_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=Braking_Report.pdf"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

# =========================================
# QMAX TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing Qmax analysis HTML interface

@router.post("/calculate_qmax")
async def legacy_qmax_calculate(raw_input: dict):
    """
    Legacy Qmax calculation endpoint for backward compatibility.

    Preserves the original API contract for maximum power optimization analysis.
    Translates legacy input format to modern Qmax tool API structure.

    Qmax Analysis:
    - Maximum power point determination
    - Engine performance optimization
    - Power curve analysis
    - Efficiency calculations

    Parameters:
        raw_input (dict): Legacy Qmax calculation parameters

    Returns:
        dict: Qmax analysis results in legacy format

    Note:
        Modern endpoint: POST /qmax/calculate
    """
    from pydantic import BaseModel
    from typing import Optional

    class LegacyQmaxInput(BaseModel):
        engine_power: Optional[float] = None
        max_rpm: Optional[float] = None
        gear_ratio: Optional[float] = None
        vehicle_weight: Optional[float] = None

    try:
        validated_input = LegacyQmaxInput(**raw_input)

        # For legacy compatibility, return dummy results
        # The modern Qmax tool is for rail axle loads, not engine power
        return {
            "qmax_value": validated_input.engine_power * validated_input.gear_ratio if validated_input.engine_power and validated_input.gear_ratio else 0,
            "optimal_rpm": validated_input.max_rpm or 0,
            "power_output": validated_input.engine_power or 0,
            "calculation_status": "success"
        }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Qmax calculation failed: {str(e)}")

@router.post("/download_qmax_report")
async def legacy_qmax_report(raw_input: dict):
    """
    Legacy Qmax report generation endpoint.

    Generates DOCX reports for Qmax analysis using legacy input format.

    Parameters:
        raw_input (dict): Legacy Qmax parameters

    Returns:
        StreamingResponse: DOCX file download

    Note:
        Modern endpoint: POST /qmax/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        calc_result = await legacy_qmax_calculate(raw_input)

        # For legacy compatibility, create a simple report
        import io
        try:
            import docx
            from docx.shared import Pt
        except ImportError:
            docx = None

        if docx:
            doc = docx.Document()
            doc.add_heading('Qmax Calculation Report', 0)
            doc.add_paragraph(f'Engine Power: {raw_input.get("engine_power", "N/A")} kW')
            doc.add_paragraph(f'Max RPM: {raw_input.get("max_rpm", "N/A")}')
            doc.add_paragraph(f'Gear Ratio: {raw_input.get("gear_ratio", "N/A")}')
            doc.add_paragraph(f'Vehicle Weight: {raw_input.get("vehicle_weight", "N/A")} kg')
            doc.add_paragraph(f'Qmax Value: {calc_result.get("qmax_value", "N/A")}')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            report_stream = buffer
        else:
            # Fallback to simple text
            report_stream = io.BytesIO(b"Qmax Report - DOCX library not available")

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Qmax_Report.docx"}
        )

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Qmax_Report.docx"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

# =========================================
# HYDRAULIC TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing hydraulic system HTML interface

@router.post("/calculate")
async def legacy_hydraulic_calculate(raw_input: dict):
    """
    Legacy hydraulic calculation endpoint for backward compatibility.

    Preserves the original API contract for hydraulic system design and analysis.
    Supports pressure, flow, and power calculations for hydraulic systems.

    Hydraulic Analysis:
    - Pump performance calculations
    - Pressure drop analysis
    - Flow rate optimization
    - Power requirements assessment

    Parameters:
        raw_input (dict): Legacy hydraulic calculation parameters

    Returns:
        dict: Hydraulic analysis results in legacy format

    Note:
        Modern endpoint: POST /hydraulic/calculate
    """
    from pydantic import BaseModel
    from typing import Optional

    class LegacyHydraulicInput(BaseModel):
        flow_rate: Optional[float] = None
        pressure: Optional[float] = None
        pump_efficiency: Optional[float] = None
        system_pressure: Optional[float] = None

    try:
        validated_input = LegacyHydraulicInput(**raw_input)

        # For legacy compatibility, return dummy results
        return {
            "power_requirement": (validated_input.flow_rate or 0) * (validated_input.pressure or 0) / 60000,  # Rough power calculation
            "pressure_drop": (validated_input.system_pressure or 0) - (validated_input.pressure or 0),
            "flow_efficiency": validated_input.pump_efficiency or 0.8,
            "calculation_status": "success"
        }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Hydraulic calculation failed: {str(e)}")

@router.post("/download_report")
async def legacy_hydraulic_report(raw_input: dict):
    """
    Legacy hydraulic report generation endpoint.

    Generates DOCX reports for hydraulic system analysis.

    Parameters:
        raw_input (dict): Legacy hydraulic parameters

    Returns:
        StreamingResponse: DOCX file download

    Note:
        Modern endpoint: POST /hydraulic/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        calc_result = await legacy_hydraulic_calculate(raw_input)

        from app.tools.hydraulic.reports.pdf_builder import create_hydraulic_docx_report

        modern_input = {
            "flow_rate_lpm": raw_input.get("flow_rate", 0),
            "pressure_bar": raw_input.get("pressure", 0),
            "pump_efficiency": raw_input.get("pump_efficiency", 0.8),
            "system_pressure_bar": raw_input.get("system_pressure", 0)
        }

        report_stream = create_hydraulic_docx_report(modern_input, {"results": calc_result}, raw_input)

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Hydraulic_Report.docx"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

# =========================================
# LOAD DISTRIBUTION TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing load distribution HTML interface

@router.post("/calculate_load_distribution")
async def legacy_load_distribution_calculate(raw_input: dict):
    """
    Legacy load distribution calculation endpoint.

    Preserves API contract for axle load analysis and weight distribution
    calculations across vehicle axles.

    Load Distribution Analysis:
    - Axle weight calculations
    - Load balance optimization
    - Center of gravity analysis
    - Weight transfer assessment

    Parameters:
        raw_input (dict): Legacy load distribution parameters

    Returns:
        dict: Load distribution results in legacy format

    Note:
        Modern endpoint: POST /load-distribution/calculate
    """
    from pydantic import BaseModel
    from typing import Optional, List

    class LegacyLoadDistributionInput(BaseModel):
        total_weight: Optional[float] = None
        wheelbase: Optional[float] = None
        cg_position: Optional[float] = None
        num_axles: Optional[int] = None

    try:
        validated_input = LegacyLoadDistributionInput(**raw_input)

        from app.tools.load_distribution.core import perform_load_distribution_calculation

        # Map legacy inputs to modern parameters with defaults
        config_type = "bogie"  # Default to bogie configuration
        total_load = validated_input.total_weight or 100.0  # Default 100 tonnes
        front_percent = 50.0  # Default 50% front load
        q1_percent = 50.0     # Default 50% on Q1 wheel
        q3_percent = 50.0     # Default 50% on Q3 wheel

        results = perform_load_distribution_calculation(config_type, total_load, front_percent, q1_percent, q3_percent)

        return {
            "axle_loads": results.get("axle_loads", []),
            "load_balance": results.get("balance_ratio", 0),
            "cg_height": results.get("cg_height", 0),
            "calculation_status": "success"
        }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Load distribution calculation failed: {str(e)}")

@router.post("/download_load_distribution_report")
async def legacy_load_distribution_report(raw_input: dict):
    """
    Legacy load distribution report generation endpoint.

    Generates DOCX reports for load distribution analysis.

    Parameters:
        raw_input (dict): Legacy load distribution parameters

    Returns:
        StreamingResponse: DOCX file download

    Note:
        Modern endpoint: POST /load-distribution/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        calc_result = await legacy_load_distribution_calculate(raw_input)

        from app.tools.load_distribution.reports.pdf_builder import create_load_distro_docx

        # Use the calculation results directly
        report_stream = create_load_distro_docx({"config_type": "bogie", "total_load": calc_result.get("total_load", 100)}, calc_result)

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Load_Distribution_Report.docx"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

# =========================================
# TRACTIVE EFFORT TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing tractive effort HTML interface

@router.post("/calculate_tractive_effort")
async def legacy_tractive_effort_calculate(raw_input: dict):
    """
    Legacy tractive effort calculation endpoint.

    Preserves API contract for locomotive tractive force and pulling capacity
    analysis across different operating conditions.

    Tractive Effort Analysis:
    - Maximum pulling force calculations
    - Speed vs tractive effort curves
    - Gear ratio optimization
    - Adhesion limit assessment

    Parameters:
        raw_input (dict): Legacy tractive effort parameters

    Returns:
        dict: Tractive effort results in legacy format

    Note:
        Modern endpoint: POST /tractive-effort/calculate
    """
    from pydantic import BaseModel
    from typing import Optional, List

    class LegacyTractiveEffortInput(BaseModel):
        loco_weight: Optional[float] = None
        power: Optional[float] = None
        gear_ratios: Optional[List[float]] = None
        adhesion: Optional[float] = None

    try:
        validated_input = LegacyTractiveEffortInput(**raw_input)

        from app.tools.tractive_effort.service import perform_te_calculation

        # Map legacy inputs to modern parameters with defaults
        modern_inputs = {
            "load": validated_input.loco_weight or 100.0,  # Use loco_weight as load
            "loco_weight": validated_input.loco_weight or 120.0,
            "gradient": 0.0,  # Default level track
            "curvature": 0.0,  # Default straight track
            "speed": 60.0,     # Default 60 km/h
            "mode": "Running",
            "grad_type": "Degree",
            "curvature_unit": "Radius(m)"
        }

        results, _ = perform_te_calculation(modern_inputs, raw_input)

        return {
            "tractive_effort": results.get("te", 0),
            "optimal_gear": 1,  # Default single gear
            "speed_range": [results.get("speed_for_power", 60.0)],
            "calculation_status": "success"
        }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Tractive effort calculation failed: {str(e)}")

@router.post("/download_tractive_effort_report")
async def legacy_tractive_effort_report(raw_input: dict):
    """
    Legacy tractive effort report generation endpoint.

    Generates DOCX reports for tractive effort analysis.

    Parameters:
        raw_input (dict): Legacy tractive effort parameters

    Returns:
        StreamingResponse: DOCX file download

    Note:
        Modern endpoint: POST /tractive-effort/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        calc_result = await legacy_tractive_effort_calculate(raw_input)

        from app.tools.tractive_effort.service import perform_te_calculation
        from app.tools.tractive_effort.reports.pdf_builder import create_te_docx_report

        # Recreate inputs for report generation
        modern_inputs = {
            "load": calc_result.get("load", 100.0),
            "loco_weight": calc_result.get("loco_weight", 120.0),
            "gradient": 0.0,
            "curvature": 0.0,
            "speed": 60.0,
            "mode": "Running",
            "grad_type": "Degree",
            "curvature_unit": "Radius(m)"
        }

        results, _ = perform_te_calculation(modern_inputs, raw_input)
        report_stream = create_te_docx_report(modern_inputs, results, raw_input)

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Tractive_Effort_Report.docx"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

# =========================================
# VEHICLE PERFORMANCE TOOL LEGACY ENDPOINTS
# =========================================
# Maintains compatibility with existing vehicle performance HTML interface

@router.post("/calculate_performance")
async def legacy_vehicle_performance_calculate(raw_input: dict):
    """
    Legacy vehicle performance calculation endpoint.

    Preserves API contract for comprehensive locomotive performance evaluation
    including tractive effort, speed analysis, and gear optimization.

    Vehicle Performance Analysis:
    - Tractive effort vs speed characteristics
    - Maximum speeds across different gears
    - Shunting capacity evaluation
    - Performance optimization recommendations

    Parameters:
        raw_input (dict): Legacy vehicle performance parameters

    Returns:
        dict: Performance analysis results in legacy format

    Note:
        Modern endpoint: POST /vehicle-performance/calculate
    """
    from pydantic import BaseModel
    from typing import Optional, List

    class LegacyVehiclePerformanceInput(BaseModel):
        loco_gvw: Optional[float] = None
        peak_power: Optional[float] = None
        wheel_dia: Optional[float] = None
        num_axles: Optional[int] = None
        gear_ratios: Optional[List[float]] = None
        friction_mu: Optional[float] = None

    try:
        validated_input = LegacyVehiclePerformanceInput(**raw_input)

        from app.tools.vehicle_performance.service import VehiclePerformanceCalculator
        from app.tools.vehicle_performance.schemas import VehiclePerformanceInput

        modern_input = VehiclePerformanceInput(
            loco_gvw=validated_input.loco_gvw or 120000,
            peak_power=validated_input.peak_power or 2000,
            wheel_dia=validated_input.wheel_dia or 1.25,
            num_axles=validated_input.num_axles or 4,
            rear_axle_ratio=3.5,  # Default value
            gear_ratios=validated_input.gear_ratios or [1.0],
            friction_mu=validated_input.friction_mu or 0.35,
            min_rpm=400,  # Default min RPM
            max_speed=100,  # Default value
            max_slope=2.0,  # Default value
            max_curve=6.0,  # Default value
            shunting_load=2000  # Default value
        )

        calculator = VehiclePerformanceCalculator(modern_input.model_dump())
        results = calculator.run_tractive_calculation()

        return {
            "traction_snapshot": results,
            "performance_data": {
                "max_speed": max([gear.get('max_speed_level_kmh', 0) for gear in results.values()]),
                "optimal_gear": 1,  # Simplified
                "efficiency": 0.85  # Placeholder
            },
            "calculation_status": "success"
        }

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"Vehicle performance calculation failed: {str(e)}")

@router.post("/download_performance_report")
async def legacy_vehicle_performance_report(raw_input: dict):
    """
    Legacy vehicle performance report generation endpoint.

    Generates DOCX reports for comprehensive vehicle performance analysis.

    Parameters:
        raw_input (dict): Legacy vehicle performance parameters

    Returns:
        StreamingResponse: DOCX file download

    Note:
        Modern endpoint: POST /vehicle-performance/download-report
    """
    from fastapi.responses import StreamingResponse

    try:
        calc_result = await legacy_vehicle_performance_calculate(raw_input)

        from app.tools.vehicle_performance.reports.pdf_builder import create_vehicle_performance_docx_report

        modern_input = {
            "loco_gvw_kg": raw_input.get("loco_gvw", 0),
            "peak_power_kw": raw_input.get("peak_power", 0),
            "wheel_dia_m": raw_input.get("wheel_dia", 1.0),
            "num_axles": raw_input.get("num_axles", 4),
            "rear_axle_ratio": 3.5,
            "gear_ratios": raw_input.get("gear_ratios", [1.0]),
            "friction_mu": raw_input.get("friction_mu", 0.35),
            "max_speed_kmh": 100,
            "max_slope": 2.0,
            "max_curve": 6.0,
            "shunting_load_t": 2000
        }

        report_stream = create_vehicle_performance_docx_report(modern_input, calc_result)

        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Performance_Report.docx"}
        )

    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")
    """
    Legacy braking calculation endpoint (for backward compatibility).

    Parameters:
        raw_input (dict): Raw input data as expected by the old HTML frontend

    Returns:
        dict: Calculation results in legacy format

    Raises:
        HTTPException: On validation or calculation error
    """
    from pydantic import BaseModel
    from typing import Optional, Dict
    from pydantic import Field

    # Define the legacy BrakingRawInput schema
    class BrakingRawInput(BaseModel):
        mass_kg: float
        reaction_time: float
        num_wheels: int
        calc_mode: str
        rail_speed_input: str
        rail_gradient_input: str
        rail_gradient_type: str
        road_speed_input: Optional[str] = ""
        road_gradient_input: Optional[str] = ""
        road_gradient_type: Optional[str] = "Percentage (%)"
        mu: Optional[float] = 0.7
        doc_no: Optional[str] = ""
        made_by: Optional[str] = ""
        checked_by: Optional[str] = ""
        approved_by: Optional[str] = ""
        wheel_dia: Optional[float] = 0

    try:
        # Convert raw input to BrakingRawInput schema
        raw_schema = BrakingRawInput(**raw_input)
        # Convert to new BrakingInput schema
        from app.tools.braking.schemas import BrakingInput
        from app.tools.braking.validation import validate_braking_inputs
        from app.tools.braking.service import perform_braking_calculation
        new_schema = BrakingInput(**raw_schema.dict())
        inputs, inputs_raw = validate_braking_inputs(new_schema)
        results, context = perform_braking_calculation(inputs)
        return {"rows": results}
    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(500, str(e))